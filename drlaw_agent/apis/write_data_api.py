import requests
from drlaw_agent.apis.base import DOMAIN, headers, record_call
import docx
import io


@record_call
def save_dict_list_to_word(company_name: str, dict_list: str) -> str:
    """
    通过传入结构化信息，制作生成公司数据报告
    例如：
        输入：
        {
            "company_name": "北京碧水源科技股份有限公司",
            "dict_list": "{}"
        }
        输出: 文件字节流
    """

    url = f"https://{DOMAIN}/law_api/s1_b/save_dict_list_to_word"

    data = {"company_name": company_name, "dict_list": dict_list}

    rsp = requests.post(url, json=data, headers=headers)
    rsp_obj = rsp.json()

    if isinstance(rsp_obj, dict):
        return [rsp_obj]
    return rsp_obj

    # # 使用 io.BytesIO 创建一个内存文件
    # memory_file = io.BytesIO(rsp.content)
    # # 使用 python-docx 从内存文件中读取文档
    # doc = docx.Document(memory_file)

    # # 提取文本
    # # 获取文档中的所有段落和表格
    # paragraphs = []
    # for para in doc.paragraphs:
    #     paragraphs.append(para.text + "\n")
    # tables = []
    # for table in doc.tables:
    #     table_str = []
    #     for row in table.rows:
    #         row_content = [cell.text for cell in row.cells]
    #         table_str.append(" | ".join(row_content) + "\n")
    #     tables.append("".join(table_str))
    # # 将提取的文本合并为一个字符串
    # text_string = ""

    # assert len(paragraphs) == len(tables)
    # for i in range(len(paragraphs)):
    #     text_string += paragraphs[i] + tables[i] + "\n"

    # return text_string


@record_call
def get_citizens_sue_citizens(info: dict) -> str:
    """
    民事起诉状(公民起诉公民)

    例如：
        输入：
        {
            "原告": "张三",
            "原告性别": "男",
            "原告生日": "1976-10-2",
            "原告民族": "汉",
            "原告工作单位": "XXX",
            "原告地址": "中国",
            "原告联系方式": "123456",
            "原告委托诉讼代理人": "李四",
            "原告委托诉讼代理人联系方式": "421313",
            "被告": "王五",
            "被告性别": "女",
            "被告生日": "1975-02-12",
            "被告民族": "汉",
            "被告工作单位": "YYY",
            "被告地址": "江苏",
            "被告联系方式": "56354321",
            "被告委托诉讼代理人": "赵六",
            "被告委托诉讼代理人联系方式": "09765213",
            "诉讼请求": "AA纠纷",
            "事实和理由": "上诉",
            "证据": "PPPPP",
            "法院名称": "最高法",
            "起诉日期": "2012-09-08"
        }

        输出:起诉状
    """

    url = f"https://{DOMAIN}/law_api/s1_b/get_citizens_sue_citizens"

    data = info

    rsp = requests.post(url, json=data, headers=headers)
    rsp_obj = rsp.json()

    if isinstance(rsp_obj, dict):
        return [rsp_obj]
    return rsp_obj


@record_call
def get_company_sue_citizens(info: dict) -> str:
    """
    民事起诉状(公司起诉公民)

    例如：
        输入：
        {
            "原告": "上海公司",
            "原告地址": "上海",
            "原告法定代表人": "张三",
            "原告联系方式": "872638",
            "原告委托诉讼代理人": "B律师事务所",
            "原告委托诉讼代理人联系方式": "5678900",
            "被告": "王五",
            "被告性别": "女",
            "被告生日": "1975-02-12",
            "被告民族": "汉",
            "被告工作单位": "YYY",
            "被告地址": "江苏",
            "被告联系方式": "56354321",
            "被告委托诉讼代理人": "赵六",
            "被告委托诉讼代理人联系方式": "09765213",
            "诉讼请求": "AA纠纷",
            "事实和理由": "上诉",
            "证据": "PPPPP",
            "法院名称": "最高法",
            "起诉日期": "2012-09-08"
        }

        输出:起诉状
    """

    url = f"https://{DOMAIN}/law_api/s1_b/get_company_sue_citizens"

    data = info

    rsp = requests.post(url, json=data, headers=headers)
    rsp_obj = rsp.json()

    if isinstance(rsp_obj, dict):
        return [rsp_obj]
    return rsp_obj


@record_call
def get_citizens_sue_company(info: dict) -> str:
    """
    民事起诉状(公民起诉公司)

    例如：
        输入：
        {
            "原告": "张三",
            "原告性别": "男",
            "原告生日": "1976-10-2",
            "原告民族": "汉",
            "原告工作单位": "XXX",
            "原告地址": "中国",
            "原告联系方式": "123456",
            "原告委托诉讼代理人": "李四",
            "原告委托诉讼代理人联系方式": "421313",
            "被告": "王五公司",
            "被告地址": "公司地址",
            "被告法定代表人": "赵四",
            "被告联系方式": "98766543",
            "被告委托诉讼代理人": "C律师事务所",
            "被告委托诉讼代理人联系方式": "425673398",
            "诉讼请求": "AA纠纷",
            "事实和理由": "上诉",
            "证据": "PPPPP",
            "法院名称": "最高法",
            "起诉日期": "2012-09-08"
        }

        输出:起诉状
    """

    url = f"https://{DOMAIN}/law_api/s1_b/get_citizens_sue_company"

    data = info

    rsp = requests.post(url, json=data, headers=headers)
    rsp_obj = rsp.json()

    if isinstance(rsp_obj, dict):
        return [rsp_obj]
    return rsp_obj


@record_call
def get_company_sue_company(info: dict) -> str:
    """
    民事起诉状(公司起诉公司)

    例如：
        输入：
        {
            "原告": "上海公司",
            "原告地址": "上海",
            "原告法定代表人": "张三",
            "原告联系方式": "872638",
            "原告委托诉讼代理人": "B律师事务所",
            "原告委托诉讼代理人联系方式": "5678900",
            "被告": "王五公司",
            "被告地址": "公司地址",
            "被告法定代表人": "赵四",
            "被告联系方式": "98766543",
            "被告委托诉讼代理人": "C律师事务所",
            "被告委托诉讼代理人联系方式": "425673398",
            "诉讼请求": "AA纠纷",
            "事实和理由": "上诉",
            "证据": "PPPPP",
            "法院名称": "最高法",
            "起诉日期": "2012-09-08"
        }

        输出:起诉状
    """

    url = f"https://{DOMAIN}/law_api/s1_b/get_company_sue_company"

    data = info

    rsp = requests.post(url, json=data, headers=headers)
    rsp_obj = rsp.json()

    if isinstance(rsp_obj, dict):
        return [rsp_obj]
    return rsp_obj


if __name__ == "__main__":
    # print(
    #     get_citizens_sue_citizens(
    #         {
    #             "原告": "张三",
    #             "原告性别": "男",
    #             "原告生日": "1976-10-2",
    #             "原告民族": "汉",
    #             "原告工作单位": "XXX",
    #             "原告地址": "中国",
    #             "原告联系方式": "123456",
    #             "原告委托诉讼代理人": "李四",
    #             "原告委托诉讼代理人联系方式": "421313",
    #             "被告": "王五",
    #             "被告性别": "女",
    #             "被告生日": "1975-02-12",
    #             "被告民族": "汉",
    #             "被告工作单位": "YYY",
    #             "被告地址": "江苏",
    #             "被告联系方式": "56354321",
    #             "被告委托诉讼代理人": "赵六",
    #             "被告委托诉讼代理人联系方式": "09765213",
    #             "诉讼请求": "AA纠纷",
    #             "事实和理由": "上诉",
    #             "证据": "PPPPP",
    #             "法院名称": "最高法",
    #             "起诉日期": "2012-09-08",
    #         }
    #     )
    # )
    # print(
    #     get_company_sue_citizens(
    #         {
    #             "原告": "上海公司",
    #             "原告地址": "上海",
    #             "原告法定代表人": "张三",
    #             "原告联系方式": "872638",
    #             "原告委托诉讼代理人": "B律师事务所",
    #             "原告委托诉讼代理人联系方式": "5678900",
    #             "被告": "王五",
    #             "被告性别": "女",
    #             "被告生日": "1975-02-12",
    #             "被告民族": "汉",
    #             "被告工作单位": "YYY",
    #             "被告地址": "江苏",
    #             "被告联系方式": "56354321",
    #             "被告委托诉讼代理人": "赵六",
    #             "被告委托诉讼代理人联系方式": "09765213",
    #             "诉讼请求": "AA纠纷",
    #             "事实和理由": "上诉",
    #             "证据": "PPPPP",
    #             "法院名称": "最高法",
    #             "起诉日期": "2012-09-08",
    #         }
    #     )
    # )
    # print(
    #     get_citizens_sue_company(
    #         {
    #             "原告": "张三",
    #             "原告性别": "男",
    #             "原告生日": "1976-10-2",
    #             "原告民族": "汉",
    #             "原告工作单位": "XXX",
    #             "原告地址": "中国",
    #             "原告联系方式": "123456",
    #             "原告委托诉讼代理人": "李四",
    #             "原告委托诉讼代理人联系方式": "421313",
    #             "被告": "王五公司",
    #             "被告地址": "公司地址",
    #             "被告法定代表人": "赵四",
    #             "被告联系方式": "98766543",
    #             "被告委托诉讼代理人": "C律师事务所",
    #             "被告委托诉讼代理人联系方式": "425673398",
    #             "诉讼请求": "AA纠纷",
    #             "事实和理由": "上诉",
    #             "证据": "PPPPP",
    #             "法院名称": "最高法",
    #             "起诉日期": "2012-09-08",
    #         }
    #     )
    # )

    # print(
    #     get_company_sue_company(
    #         {
    #             "原告": "上海公司",
    #             "原告地址": "上海",
    #             "原告法定代表人": "张三",
    #             "原告联系方式": "872638",
    #             "原告委托诉讼代理人": "B律师事务所",
    #             "原告委托诉讼代理人联系方式": "5678900",
    #             "被告": "王五公司",
    #             "被告地址": "公司地址",
    #             "被告法定代表人": "赵四",
    #             "被告联系方式": "98766543",
    #             "被告委托诉讼代理人": "C律师事务所",
    #             "被告委托诉讼代理人联系方式": "425673398",
    #             "诉讼请求": "AA纠纷",
    #             "事实和理由": "上诉",
    #             "证据": "PPPPP",
    #             "法院名称": "最高法",
    #             "起诉日期": "2012-09-08",
    #         }
    #     )
    # )

    print(
        save_dict_list_to_word(
            company_name="北京碧水源科技股份有限公司",
            dict_list="{'工商信息': [{'公司名称': '北京碧水源科技股份有限公司', '登记状态': '存续', '统一社会信用代码': '91110000802115985Y', '参保人数': '351', '行业一级': '科学研究和技术服务业', '行业二级': '科技推广和应用服务业', '行业三级': '其他科技推广服务业'}], '子公司信息': [{'关联上市公司全称': '北京碧水源科技股份有限公司', '上市公司关系': '子公司', '上市公司参股比例': 100.0, '上市公司投资金额': '1.06亿', '公司名称': '北京碧海环境科技有限公司'}], '裁判文书': [{'关联公司': '北京碧水源科技股份有限公司', '原告': '四川帝宇水利水电工程有限公司', '被告': '成都碧水源江环保科技有限公司,北京碧水源科技股份有限公司', '案由': '建设工程施工合同纠纷', '涉案金额': 0.0, '日期': Timestamp('2019-07-23 00:00:00')}], '限制高消费': [{'限制高消费企业名称': '南京仙林碧水源污水处理有限公司', '案号': '（2024）苏0113执1601号', '申请人': '苏华建设集团有限公司', '涉案金额': '-', '立案日期': Timestamp('2024-04-07 00:00:00'), '限高发布日期': Timestamp('2024-06-24 00:00:00')}]}",
        )
    )
